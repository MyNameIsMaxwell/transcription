import os
import uuid
import shutil
import asyncio
from concurrent.futures import ThreadPoolExecutor
from typing import Dict, Optional
from enum import Enum
from fastapi import FastAPI, File, UploadFile, HTTPException, Request, Form
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from starlette.middleware.sessions import SessionMiddleware
from pydub import AudioSegment
import openai
from openai import OpenAI

from dotenv import load_dotenv
from docx import Document
from faster_whisper import WhisperModel
import subprocess
import base64
import json

load_dotenv()
client = OpenAI(api_key=os.getenv("DEEPSEEK_KEY"), base_url="https://api.deepseek.com")

app = FastAPI()
# Используем middleware для управления сессиями
app.add_middleware(SessionMiddleware, secret_key="!secret")

# Определяем директории для хранения загруженных файлов и результатов
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_DIR = os.path.join(BASE_DIR, "uploads")
RESULT_DIR = os.path.join(BASE_DIR, "results")
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(RESULT_DIR, exist_ok=True)

# ThreadPoolExecutor для фоновых задач (чтобы не блокировать event loop)
executor = ThreadPoolExecutor(max_workers=2)

# Хранение статусов задач
task_statuses: Dict[str, Dict] = {}

class TranscriptionEngine(str, Enum):
    WHISPER = "whisper"
    SPEECHKIT = "speechkit"

# Кэш для модели Whisper (загружается один раз)
_whisper_model: Optional[WhisperModel] = None

def get_whisper_model() -> WhisperModel:
    """Получает или создает модель Whisper (кэшируется)"""
    global _whisper_model
    if _whisper_model is None:
        model_size = os.getenv("WHISPER_MODEL_SIZE", "base")
        compute_type = os.getenv("WHISPER_COMPUTE_TYPE", "int8")
        num_threads = int(os.getenv("WHISPER_NUM_THREADS", "4"))
        _whisper_model = WhisperModel(
            model_size,
            device="cpu",
            compute_type=compute_type,
            num_workers=num_threads
        )
    return _whisper_model

def check_ffmpeg():
    """Проверяет наличие ffmpeg в системе"""
    ffmpeg_path = os.getenv("FFMPEG_PATH", "ffmpeg")
    try:
        result = subprocess.run(
            [ffmpeg_path, "-version"],
            capture_output=True,
            timeout=5
        )
        if result.returncode == 0:
            return True
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    
    # Пробуем найти в стандартных местах
    if os.path.exists(ffmpeg_path):
        return True
    
    raise RuntimeError(
        f"FFmpeg не найден. Установите FFmpeg и добавьте его в PATH, "
        f"или укажите путь через переменную окружения FFMPEG_PATH. "
        f"Текущий путь: {ffmpeg_path}"
    )

# Проверяем ffmpeg при старте приложения
try:
    check_ffmpeg()
except RuntimeError as e:
    print(f"Предупреждение: {e}")

# Маршрут для главной страницы (фронтенд)
@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    # Если в сессии отсутствует user_id, генерируем новый
    user_id = request.session.get("user_id")
    if not user_id:
        user_id = str(uuid.uuid4())
        request.session["user_id"] = user_id
    # Читаем файл index.html (находится в той же папке)
    index_path = os.path.join(BASE_DIR, "index.html")
    with open(index_path, "r", encoding="utf-8") as f:
        html_content = f.read()
    return HTMLResponse(content=html_content, media_type="text/html")

# Эндпоинт загрузки аудиофайла. Можно передать дополнительный параметр "prompt" и "engine"
@app.post("/upload/")
async def upload_audio(
    request: Request,
    file: UploadFile = File(...),
    prompt: str = Form(None),
    engine: str = Form("whisper")
):
    user_id = request.session.get("user_id")
    if not user_id:
        user_id = str(uuid.uuid4())
        request.session["user_id"] = user_id

    # Валидация движка
    try:
        engine_enum = TranscriptionEngine(engine.lower())
    except ValueError:
        raise HTTPException(status_code=400, detail=f"Неизвестный движок: {engine}. Доступны: whisper, speechkit")

    # Сохраняем загруженный файл на сервер
    file_id = str(uuid.uuid4())
    filename = f"{file_id}_{file.filename}"
    file_path = os.path.join(UPLOAD_DIR, filename)
    with open(file_path, "wb") as f_out:
        content = await file.read()
        f_out.write(content)

    # Если пользователь не задал промпт, используем значение по умолчанию
    if not prompt:
        prompt = "Создай краткое содержание в виде списка ключевых тезисов на основе текста пользователя"

    # Инициализируем статус задачи
    task_statuses[file_id] = {
        "status": "queued",
        "progress": 0,
        "message": "Задача поставлена в очередь",
        "download_url": None,
        "error": None
    }

    # Запускаем фоновую задачу для обработки аудио через executor
    loop = asyncio.get_event_loop()
    loop.run_in_executor(executor, process_audio, file_path, file_id, user_id, prompt, engine_enum)

    return JSONResponse({
        "message": "Файл загружен, обработка запущена",
        "file_id": file_id
    })

# Эндпоинт для получения статуса задачи
@app.get("/status/{file_id}")
async def get_status(file_id: str):
    if file_id not in task_statuses:
        raise HTTPException(status_code=404, detail="Задача не найдена")
    return JSONResponse(task_statuses[file_id])

def process_audio(file_path: str, file_id: str, user_id: str, prompt: str, engine: TranscriptionEngine):
    try:
        # Обновляем статус: начата обработка
        task_statuses[file_id].update({
            "status": "running",
            "progress": 10,
            "message": "Начата обработка аудио"
        })

        # 1. Разбивка аудио на части, если длительность больше 10 минут
        task_statuses[file_id].update({
            "progress": 20,
            "message": "Разбивка аудио на части..."
        })
        parts = split_audio(file_path)
        
        # 2. Транскрибация
        task_statuses[file_id].update({
            "progress": 30,
            "message": f"Транскрибация через {engine.value}..."
        })
        transcripts = []
        total_parts = len(parts)
        for idx, part in enumerate(parts):
            if engine == TranscriptionEngine.WHISPER:
                transcript = transcribe_with_whisper(part)
            elif engine == TranscriptionEngine.SPEECHKIT:
                transcript = transcribe_with_speechkit(part)
            else:
                raise ValueError(f"Неизвестный движок: {engine}")
            
            transcripts.append(transcript)
            # Если файл является временным (от разбиения), удаляем его после транскрипции
            if part != file_path:
                os.remove(part)
            
            # Обновляем прогресс
            progress = 30 + int((idx + 1) / total_parts * 50)
            task_statuses[file_id].update({
                "progress": progress,
                "message": f"Транскрибация: {idx + 1}/{total_parts} частей"
            })
        
        full_text = " ".join(transcripts).strip()

        # Проверяем, что транскрипция не пуста
        if not full_text:
            raise ValueError("Транскрипция пуста. Возможно, аудио не содержит речи или произошла ошибка распознавания.")

        # 3. Генерация резюме через DeepSeek API
        task_statuses[file_id].update({
            "progress": 85,
            "message": "Генерация резюме..."
        })
        summary = generate_summary(full_text, prompt)

        # 4. Создание итогового Word файла (.docx)
        task_statuses[file_id].update({
            "progress": 90,
            "message": "Создание документа..."
        })
        result_filename = create_docx(summary, file_id)

        # 5. Сохранение результата в истории пользователя (директория RESULT_DIR/user_id)
        user_dir = os.path.join(RESULT_DIR, user_id)
        os.makedirs(user_dir, exist_ok=True)
        final_path = os.path.join(user_dir, result_filename)
        # Используем абсолютный путь для move
        abs_result_path = os.path.join(BASE_DIR, result_filename)
        if os.path.exists(abs_result_path):
            shutil.move(abs_result_path, final_path)
        elif os.path.exists(result_filename):
            shutil.move(result_filename, final_path)

        # 6. Удаляем исходный аудиофайл, если он существует
        if os.path.exists(file_path):
            os.remove(file_path)

        # Обновляем статус: завершено
        task_statuses[file_id].update({
            "status": "done",
            "progress": 100,
            "message": "Обработка завершена",
            "download_url": f"/download/{user_id}/{result_filename}"
        })
    except Exception as e:
        error_msg = str(e)
        print(f"Error processing file {file_id}: {error_msg}")
        task_statuses[file_id].update({
            "status": "error",
            "message": f"Ошибка обработки: {error_msg}",
            "error": error_msg
        })

def prepare_audio_for_transcription(audio_path: str) -> str:
    """
    Подготавливает аудио для транскрибации: конвертирует в моно, 16kHz, WAV.
    Возвращает путь к подготовленному файлу.
    """
    try:
        audio = AudioSegment.from_file(audio_path)
        # Конвертируем в моно и 16kHz для оптимизации
        audio = audio.set_channels(1)
        audio = audio.set_frame_rate(16000)
        
        # Сохраняем во временный WAV файл
        base, _ = os.path.splitext(audio_path)
        prepared_path = f"{base}_prepared.wav"
        audio.export(prepared_path, format="wav")
        return prepared_path
    except Exception as e:
        print(f"Ошибка подготовки аудио {audio_path}: {e}")
        # Если не удалось подготовить, возвращаем исходный файл
        return audio_path

def split_audio(file_path: str):
    """
    Разбивает аудиофайл на части, если его длительность превышает 10 минут.
    Возвращает список путей к файлам (если длительность меньше – возвращает исходный файл).
    """
    try:
        audio = AudioSegment.from_file(file_path)
    except Exception as e:
        raise RuntimeError(
            f"Не удалось загрузить аудиофайл {file_path}. "
            f"Убедитесь, что файл не поврежден и FFmpeg установлен корректно. Ошибка: {str(e)}"
        )
    
    duration_ms = len(audio)
    max_duration = 10 * 60 * 1000  # 10 минут в миллисекундах

    if duration_ms <= max_duration:
        return [file_path]

    parts = []
    num_parts = (duration_ms // max_duration) + (1 if duration_ms % max_duration else 0)
    base, ext = os.path.splitext(file_path)
    for i in range(num_parts):
        try:
            start_ms = i * max_duration
            end_ms = min((i + 1) * max_duration, duration_ms)
            chunk = audio[start_ms:end_ms]
            part_filename = f"{base}_part{i+1}{ext}"
            # Экспортируем часть; формат определяется по расширению (без точки)
            export_format = ext[1:] if ext else "wav"
            chunk.export(part_filename, format=export_format)
            parts.append(part_filename)
        except Exception as e:
            # Очищаем уже созданные части при ошибке
            for part in parts:
                try:
                    if os.path.exists(part):
                        os.remove(part)
                except:
                    pass
            raise RuntimeError(f"Ошибка при разбиении аудио на части: {str(e)}")
    
    return parts


def transcribe_with_whisper(audio_path: str) -> str:
    """
    Выполняет транскрипцию аудио через локальный faster-whisper (CPU).
    """
    try:
        # Подготавливаем аудио (моно, 16kHz)
        prepared_path = prepare_audio_for_transcription(audio_path)
        is_prepared = (prepared_path != audio_path)
        
        try:
            # Получаем модель (кэшируется)
            model = get_whisper_model()
            
            # Выполняем транскрипцию
            segments, info = model.transcribe(
                prepared_path,
                beam_size=5,
                language="ru",  # Можно сделать настраиваемым через env
                task="transcribe"
            )
            
            # Собираем текст из сегментов
            text_parts = []
            for segment in segments:
                text_parts.append(segment.text)
            
            result = " ".join(text_parts).strip()
            
            # Удаляем временный подготовленный файл, если он был создан
            if is_prepared and os.path.exists(prepared_path):
                os.remove(prepared_path)
            
            return result
        except Exception as e:
            # Удаляем временный файл даже при ошибке
            if is_prepared and os.path.exists(prepared_path):
                try:
                    os.remove(prepared_path)
                except:
                    pass
            raise e
    except Exception as e:
        print(f"Error transcribing {audio_path} with Whisper: {str(e)}")
        raise

def transcribe_with_speechkit(audio_path: str) -> str:
    """
    Выполняет транскрипцию аудио через Yandex SpeechKit REST API.
    Для длинных аудио разбивает на короткие фрагменты (20-30 секунд).
    """
    try:
        api_key = os.getenv("YANDEX_API_KEY")
        folder_id = os.getenv("YANDEX_FOLDER_ID")
        lang = os.getenv("YANDEX_LANG", "ru-RU")
        
        if not api_key or not folder_id:
            raise ValueError(
                "Yandex SpeechKit требует YANDEX_API_KEY и YANDEX_FOLDER_ID. "
                "Установите их в переменных окружения."
            )
        
        # Загружаем аудио
        audio = AudioSegment.from_file(audio_path)
        duration_ms = len(audio)
        
        # SpeechKit REST API имеет ограничения, поэтому разбиваем на фрагменты по 25 секунд
        chunk_duration_ms = 25 * 1000  # 25 секунд
        
        if duration_ms <= chunk_duration_ms:
            # Короткое аудио - обрабатываем целиком
            return _transcribe_speechkit_chunk(audio, api_key, folder_id, lang)
        else:
            # Длинное аудио - разбиваем на части
            transcripts = []
            num_chunks = (duration_ms // chunk_duration_ms) + (1 if duration_ms % chunk_duration_ms else 0)
            
            for i in range(num_chunks):
                start_ms = i * chunk_duration_ms
                end_ms = min((i + 1) * chunk_duration_ms, duration_ms)
                chunk = audio[start_ms:end_ms]
                
                # Транскрибируем чанк
                chunk_text = _transcribe_speechkit_chunk(chunk, api_key, folder_id, lang)
                if chunk_text:
                    transcripts.append(chunk_text)
            
            return " ".join(transcripts)
    except Exception as e:
        print(f"Error transcribing {audio_path} with SpeechKit: {str(e)}")
        raise

def _transcribe_speechkit_chunk(audio_segment: AudioSegment, api_key: str, folder_id: str, lang: str) -> str:
    """
    Транскрибирует один фрагмент аудио через SpeechKit REST API.
    """
    import requests
    
    # Конвертируем в OGG Opus (формат, поддерживаемый SpeechKit)
    # Создаем временный файл
    temp_file = os.path.join(UPLOAD_DIR, f"temp_{uuid.uuid4()}.ogg")
    try:
        # Экспортируем в OGG Opus
        audio_segment.export(temp_file, format="ogg", codec="libopus", parameters=["-ar", "16000"])
        
        # Читаем файл как бинарные данные
        with open(temp_file, "rb") as f:
            audio_data = f.read()
        
        # Формируем запрос к SpeechKit
        url = "https://stt.api.cloud.yandex.net/speech/v1/stt:recognize"
        headers = {
            "Authorization": f"Api-Key {api_key}"
        }
        
        # Параметры в query string
        params = {
            "folderId": folder_id,
            "format": "oggopus",
            "lang": lang,
            "sampleRateHertz": "16000",
            "topic": "general"  # Можно настроить через env
        }
        
        # Отправляем запрос: бинарные данные в теле, параметры в query
        response = requests.post(
            url,
            headers=headers,
            params=params,
            data=audio_data
        )
        
        if response.status_code != 200:
            error_text = response.text
            raise RuntimeError(f"SpeechKit API вернул ошибку {response.status_code}: {error_text}")
        
        result = response.json()
        
        # Извлекаем текст из ответа
        if "result" in result and "alternatives" in result["result"]:
            alternatives = result["result"]["alternatives"]
            if alternatives and len(alternatives) > 0:
                return alternatives[0].get("text", "")
        
        return ""
    finally:
        # Удаляем временный файл
        if os.path.exists(temp_file):
            try:
                os.remove(temp_file)
            except:
                pass

def generate_summary(text: str, prompt: str) -> str:
    """
    Генерирует резюме, отправляя полный текст в DeepSeek API с указанным системным промптом.
    """
    try:
        deepseek_key = os.getenv("DEEPSEEK_KEY")
        if not deepseek_key:
            raise ValueError(
                "DEEPSEEK_KEY не установлен. Установите его в переменных окружения для генерации резюме."
            )
        
        messages = [
            {"role": "system", "content": prompt},
            {"role": "user", "content": text}
        ]
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=messages,
            temperature=0.5,
            max_tokens=1000
        )
        # Обращаемся к результату через словарный синтаксис
        return response.choices[0].message.content
    except ValueError as e:
        # Пробрасываем ValueError дальше, чтобы пользователь видел понятную ошибку
        raise
    except Exception as e:
        error_msg = f"Ошибка генерации резюме через DeepSeek API: {str(e)}"
        print(error_msg)
        raise RuntimeError(error_msg)

def create_docx(summary: str, file_id: str):
    """
    Создает .docx файл с итоговым резюме.
    """
    doc = Document()
    doc.add_heading("Резюме", level=1)
    doc.add_paragraph(summary)
    output_filename = f"{file_id}_summary.docx"
    # Сохраняем в абсолютный путь
    output_path = os.path.join(BASE_DIR, output_filename)
    doc.save(output_path)
    return output_filename

# Эндпоинт для скачивания итогового файла
@app.get("/download/{user_id}/{filename}")
async def download_file(user_id: str, filename: str):
    file_path = os.path.join(RESULT_DIR, user_id, filename)
    if os.path.exists(file_path):
        return FileResponse(
            file_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=filename
        )
    raise HTTPException(status_code=404, detail="Файл не найден")

# Эндпоинт для получения истории обработанных файлов текущего пользователя
@app.get("/history", response_class=JSONResponse)
async def get_history(request: Request):
    user_id = request.session.get("user_id")
    if not user_id:
        return JSONResponse({"history": [], "user_id": None})
    user_dir = os.path.join(RESULT_DIR, user_id)
    if not os.path.exists(user_dir):
        return JSONResponse({"history": [], "user_id": user_id})
    files = os.listdir(user_dir)
    return JSONResponse({"history": files, "user_id": user_id})

# При необходимости можно смонтировать статические файлы (например, для CSS или JS)
app.mount("/static", StaticFiles(directory=os.path.join(BASE_DIR, "static")), name="static")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", port=8000, reload=True)
